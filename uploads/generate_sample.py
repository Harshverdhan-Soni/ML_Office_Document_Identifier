from reportlab.pdfgen import canvas
from docx import Document
from openpyxl import Workbook
from PIL import Image, ImageDraw

# 1. Sample Bill (PDF)
def create_pdf(filename):
    c = canvas.Canvas(filename)
    c.setFont("Helvetica", 12)
    c.drawString(100, 750, "Sample Bill")
    c.drawString(100, 730, "Bill No: 001")
    c.drawString(100, 710, "Date: 2025-08-19")
    c.drawString(100, 690, "Amount: Rs. 5000")
    c.save()

# 2. Sample Note (DOCX)
def create_note_docx(filename):
    doc = Document()
    doc.add_heading("Office Note", 0)
    doc.add_paragraph("This is a sample office note regarding administrative tasks.")
    doc.save(filename)

# 3. Sample Invoice (XLSX)
def create_invoice_xlsx(filename):
    wb = Workbook()
    ws = wb.active
    ws.title = "Invoice"
    ws.append(["Invoice No", "Date", "Amount"])
    ws.append(["INV-001", "2025-08-19", "7500"])
    wb.save(filename)

# 4. Sample Indent (DOCX)
def create_indent_docx(filename):
    doc = Document()
    doc.add_heading("Indent Request", 0)
    doc.add_paragraph("This is a request for procuring 10 laptops for the office.")
    doc.save(filename)

# 5. Sample Tender (Image PNG)
def create_tender_image(filename):
    img = Image.new("RGB", (400, 200), color=(255, 255, 255))
    d = ImageDraw.Draw(img)
    d.text((10, 10), "Tender Document\nTender ID: TND123\nDue Date: 2025-09-01", fill=(0, 0, 0))
    img.save(filename)

if __name__ == "__main__":
    create_pdf("sample_bill.pdf")
    create_note_docx("sample_note.docx")
    create_invoice_xlsx("sample_invoice.xlsx")
    create_indent_docx("sample_indent.docx")
    create_tender_image("sample_tender.png")

    print("✅ Sample files generated in current folder!")
